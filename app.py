import streamlit as st
import barcode
from barcode.writer import ImageWriter
import io
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configuración de la página Streamlit
st.set_page_config(page_title="Generador de Etiquetas Compactas", layout="centered")

st.title("Generador de Etiquetas 🏷️")
st.write("Copia y pega tu lista de productos aquí:")

# Lista de tus productos cargada por defecto
productos_default = """Americano 210 ml $25, CFAMEGR
Americano 315 ml $32, CFAMG
Expresso 90 ml $28, CFEPRG
Té de Manzanilla 315 ml $32, TMLA
Capuchino 210 ml $38, CFDCF
Capuchino 315 ml $54, CFDGG
Lechero 210 ml $38, CFDLG
Lechero 315 ml $54, CFDLC
Latte Caramel Toffe 315 ml $54, CCTFFD
Latte Arroz con Leche 315 ml $54, ARLTTE
Latte Chai 315 ml $57, CFMDL
Latte Canela 315 ml $59, LTTECANDESC
Extra Extracto de Café 18 gr $18, CEXCAF
Extra Crema Batida 15 gr $12, CREMEXT
Extra Leche Deslactosada 60 ml $6, EXTLDES
Extra Jarabe Canela 25 ml $18, EXTRJCNL"""

texto_entrada = st.text_area("Formato: Nombre, Código", value=productos_default, height=250)

if st.button("Generar Hoja de Etiquetas"):
    if texto_entrada.strip():
        doc = Document()
        
        # --- AJUSTES DE MARGENES PARA APROVECHAR LA HOJA ---
        section = doc.sections[0]
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

        lineas = [linea.strip() for linea in texto_entrada.split('\n') if "," in linea]
        
        if lineas:
            # Crear tabla (3 columnas para que quepan todas en una fila)
            num_columnas = 3
            num_filas = (len(lineas) + num_columnas - 1) // num_columnas
            tabla = doc.add_table(rows=num_filas, cols=num_columnas)
            tabla.autofit = True

            code128 = barcode.get_barcode_class('code128')

            for index, linea in enumerate(lineas):
                fila_idx = index // num_columnas
                col_idx = index % num_columnas
                celda = tabla.cell(fila_idx, col_idx)
                
                nombre_producto, codigo_barras = [p.strip() for p in linea.split(",", 1)]

                # Generar código de barras
                buffer = io.BytesIO()
                mi_codigo = code128(codigo_barras, writer=ImageWriter())
                mi_codigo.write(buffer, options={
                    'write_text': False, 
                    'module_width': 0.2, # Más delgado
                    'module_height': 8,   # Más bajo
                    'quiet_zone': 1
                })
                buffer.seek(0)

                # Formato dentro de la celda
                p = celda.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # --- NOMBRE DEL PRODUCTO EN NEGRITAS ---
                run_nombre = p.add_run(f"{nombre_producto}\n")
                run_nombre.bold = True
                run_nombre.font.size = Pt(10) # Un poco más grande para resaltar
                
                # Imagen (Tamaño reducido)
                run_img = p.add_run()
                run_img.add_picture(buffer, width=Cm(4), height=Cm(1.8))
                
                # --- CÓDIGO SIN NEGRITAS ---
                run_codigo = p.add_run(f"\n{codigo_barras}")
                run_codigo.bold = False # Texto normal para crear contraste
                run_codigo.font.size = Pt(8)

            # Guardar y descargar
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            
            st.success(f"Se acomodaron {len(lineas)} etiquetas en una cuadrícula.")
            st.download_button(
                label="📥 Descargar Word (1 Hoja)",
                data=doc_buffer,
                file_name="etiquetas_compactas.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.error("No hay datos para procesar.")
